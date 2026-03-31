"""Routes tiện ích: thời tiết, tin tức, YouTube, chạy code, fetch URL, tìm kiếm, tải file và 600+ tools."""

import os
import re
import time
import zipfile
from datetime import datetime

from flask import Blueprint, jsonify, request, send_file, session
from PyPDF2 import PdfReader

from routes.tools import (
    extract_youtube_id, fetch_url, get_news, get_weather, get_weather_json,
    run_code, run_tool, search_youtube, search_youtube_multiple,
    calculate_math, unit_converter, currency_converter, bmi_calculator,
    generate_password, qr_code_url, hash_text, base64_encode_decode,
    countdown_timer, random_number, color_picker_info, loan_calculator,
    ip_lookup, my_ip, word_counter, age_calculator, timezone_converter,
    translate_text, get_stock_info, get_crypto_price, pomodoro_plan,
    generate_lorem_ipsum, motivational_quote, define_word, ping_website,
    regex_tester, fibonacci, prime_check, binary_converter, caesar_cipher,
    check_palindrome, get_current_time, color_palette_generate,
    percentage_calc, study_plan,
    # 10 New Sophisticated Tools
    analyze_seo_url, validate_json_format, validate_email_advanced,
    analyze_code_quality, analyze_sentiment, compress_url_analyzer,
    text_similarity_compare, html_to_text_analyzer, check_url_accessibility,
    analyze_sql_query,
)

utilities_bp = Blueprint("utilities", __name__)

try:
    from duckduckgo_search import DDGS
    _DDGS_AVAILABLE = True
except ImportError:
    _DDGS_AVAILABLE = False


# ──────────────────────────────────────────
#  TOOLS PAGE
# ──────────────────────────────────────────

@utilities_bp.route("/tools")
def tools_page():
    """Serve the tools dashboard HTML page."""
    from flask import send_from_directory
    return send_from_directory("static", "tools.html")


# ──────────────────────────────────────────
#  WEATHER (FIXED)
# ──────────────────────────────────────────

@utilities_bp.route("/api/weather", methods=["GET"])
def weather_route():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    city = request.args.get("city", "Ha Noi")
    result = get_weather(city)
    return jsonify({"result": result, "city": city})


@utilities_bp.route("/api/weather-json", methods=["GET"])
def weather_json_route():
    """Weather data dạng JSON cho frontend rich display."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    city = request.args.get("city", "Ha Noi")
    data = get_weather_json(city)
    return jsonify(data)


# ──────────────────────────────────────────
#  NEWS (FIXED)
# ──────────────────────────────────────────

@utilities_bp.route("/api/news", methods=["GET"])
def news_route():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    topic = request.args.get("topic", "Việt Nam hôm nay")
    items = get_news(topic, max_results=10)
    return jsonify({"result": items, "topic": topic})


@utilities_bp.route("/api/news-json", methods=["GET"])
def news_json_route():
    """News data dạng JSON list cho frontend cards."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    topic = request.args.get("topic", "Việt Nam hôm nay")
    items = get_news(topic, max_results=12)
    return jsonify({"items": items, "topic": topic, "count": len(items)})


# ──────────────────────────────────────────
#  YOUTUBE (Enhanced - by name or URL)
# ──────────────────────────────────────────

@utilities_bp.route("/api/youtube", methods=["POST"])
def youtube_route():
    data = request.get_json()
    query = (data.get("query") or "").strip()
    if not query:
        return jsonify({"error": "Thiếu query"}), 400
    vid_id = extract_youtube_id(query)  # Try URL first
    title = query
    if not vid_id:
        vid_id, title = search_youtube(query)  # Search by name
    if vid_id:
        return jsonify({
            "video_id": vid_id,
            "title": title or query,
            "embed_url": f"https://www.youtube.com/embed/{vid_id}?autoplay=1",
        })
    return jsonify({"error": "Không tìm thấy video"}), 404


@utilities_bp.route("/api/youtube/search", methods=["POST"])
def youtube_search_multiple():
    """Tìm nhiều video YouTube."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    query = (data.get("query") or "").strip()
    if not query:
        return jsonify({"error": "Thiếu query"}), 400
    results = search_youtube_multiple(query, max_results=6)
    return jsonify({"results": results, "query": query})


# ──────────────────────────────────────────
#  CODE RUNNER
# ──────────────────────────────────────────

@utilities_bp.route("/api/run-code", methods=["POST"])
def run_code_route():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    lang = (data.get("language") or data.get("lang") or "python").strip()
    code = (data.get("code") or "").strip()
    if not code:
        return jsonify({"error": "Không có code"}), 400
    start = time.time()
    output = run_code(lang, code)
    elapsed = round(time.time() - start, 3)
    return jsonify({"output": output, "elapsed": elapsed, "language": lang})


# ──────────────────────────────────────────
#  FETCH & SEARCH
# ──────────────────────────────────────────

@utilities_bp.route("/api/fetch", methods=["POST"])
def fetch_route():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    url = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "Thiếu URL"}), 400
    content = fetch_url(url)
    return jsonify({"content": content, "url": url})


@utilities_bp.route("/api/search", methods=["POST"])
def search_route():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    query = (data.get("query") or "").strip()
    if not query:
        return jsonify({"error": "Thiếu query"}), 400
    results = []
    try:
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=8):
                results.append({
                    "title": r.get("title", ""),
                    "body": r.get("body", "")[:300],
                    "href": r.get("href", ""),
                })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    return jsonify({"results": results, "query": query})


# ──────────────────────────────────────────
#  DOWNLOAD ALL
# ──────────────────────────────────────────

@utilities_bp.route("/api/download-all", methods=["POST"])
def download_all():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    filenames = data.get("files", [])
    if not filenames:
        return jsonify({"error": "Không có file nào"}), 400
    os.makedirs("outputs", exist_ok=True)
    zip_path = os.path.join("outputs", f"files_{int(time.time())}.zip")
    with zipfile.ZipFile(zip_path, "w") as zf:
        for fn in filenames:
            fp = os.path.join("outputs", fn)
            if os.path.exists(fp):
                zf.write(fp, fn)
    return send_file(zip_path, as_attachment=True, download_name="dns_bot_files.zip")


@utilities_bp.route("/api/parse-exam-pdf", methods=["POST"])
def parse_exam_pdf():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    if "file" not in request.files:
        return jsonify({"error": "Chưa gửi file"}), 400
    file = request.files["file"]
    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Chỉ hỗ trợ PDF"}), 400

    try:
        reader = PdfReader(file)
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        text = text.strip()
        if not text:
            return jsonify({"error": "PDF rỗng hoặc không đọc được"}), 400

        # cố gắng lấy JSON từ nội dung
        import json, re

        json_text = None
        # tìm khối JSON đầu tiên
        m = re.search(r"\{[\s\S]*\}", text)
        if m:
            json_text = m.group(0)

        if not json_text:
            json_text = text

        data = json.loads(json_text)
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": f"Không parse được PDF: {str(e)}"}), 400


@utilities_bp.route("/api/testgen/export", methods=["POST"])
def testgen_export():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json() or {}
    questions = data.get("questions", [])
    title = (data.get("title") or "Bài kiểm tra").strip()
    mode = (data.get("format") or "docx").lower()

    if not questions:
        return jsonify({"error": "Không có câu hỏi"}), 400

    def difficulty(points):
        try:
            p = float(points)
        except Exception:
            p = 1
        if p <= 2:
            return "Dễ"
        if p <= 4:
            return "Trung bình"
        return "Khó"

    def weight(q):
        t = (q.get("type") or "").lower()
        p = float(q.get("points", 1) or 1)
        if "tự luận" in t:
            return round(p * 1.5, 2)
        if "trắc" in t:
            return round(p * 1.2, 2)
        return round(p, 2)

    os.makedirs("outputs", exist_ok=True)
    safe_title = re.sub(r"[^a-zA-Z0-9_]+", "_", title).strip("_") or "test"

    if mode == "docx":
        try:
            from docx import Document
        except ImportError:
            return jsonify({"error": "python-docx chưa cài. pip install python-docx"}), 500

        doc = Document()
        doc.add_heading(title, 0)
        school = data.get("school", "").strip()
        subject = data.get("subject", "").strip()
        exam_time = data.get("exam_time", "").strip()
        if school:
            doc.add_paragraph(f"Trường: {school}")
        if subject:
            doc.add_paragraph(f"Môn: {subject}")
        if exam_time:
            try:
                dt = datetime.fromisoformat(exam_time)
                doc.add_paragraph(f"Thời gian thi: {dt.strftime('%d/%m/%Y %H:%M')}")
            except Exception:
                doc.add_paragraph(f"Thời gian thi: {exam_time}")
        if data.get("description"):
            doc.add_paragraph(data.get("description", ""))

        # Ma trận đề
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "#"
        hdr_cells[1].text = "Loại"
        hdr_cells[2].text = "Câu hỏi"
        hdr_cells[3].text = "Độ khó"
        hdr_cells[4].text = "Điểm"
        hdr_cells[5].text = "Trọng số"
        for idx, q in enumerate(questions, 1):
            qtype = q.get("type") or "Không xác định"
            qtext = q.get("text", "").strip()
            qpts = float(q.get("points", 1) or 1)
            n = table.add_row().cells
            n[0].text = str(idx)
            n[1].text = str(qtype)
            n[2].text = qtext
            n[3].text = difficulty(qpts)
            n[4].text = str(qpts)
            n[5].text = str(weight(q))

        doc.add_page_break()
        doc.add_heading("Đáp án", level=1)
        for idx, q in enumerate(questions, 1):
            key = q.get("answer", "")
            doc.add_paragraph(f"Câu {idx}: {key}")

        out_path = os.path.join("outputs", f"{safe_title}_{int(time.time())}.docx")
        doc.save(out_path)
        return send_file(out_path, as_attachment=True, download_name=os.path.basename(out_path))

    elif mode == "pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            return jsonify({"error": "fpdf chưa cài. pip install fpdf"}), 500

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=12)
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, title, ln=True)
        pdf.ln(4)
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 8, data.get("description", ""))
        pdf.ln(4)
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 8, 'Ma trận đề', ln=True)
        pdf.set_font("Arial", '', 10)
        pdf.ln(2)

        # Header for 6 cols
        pdf.cell(8, 7, '#', 1)
        pdf.cell(30, 7, 'Loại', 1)
        pdf.cell(80, 7, 'Câu hỏi', 1)
        pdf.cell(22, 7, 'Độ khó', 1)
        pdf.cell(16, 7, 'Điểm', 1)
        pdf.cell(24, 7, 'Trọng số', 1)
        pdf.ln()

        for idx, q in enumerate(questions, 1):
            qtype = q.get("type") or ""
            qtext = q.get("text", "").strip().replace('\n', ' ')
            qpts = float(q.get("points", 1) or 1)
            pdf.cell(8, 7, str(idx), 1)
            pdf.cell(30, 7, qtype[:16], 1)
            pdf.cell(80, 7, qtext[:45], 1)
            pdf.cell(22, 7, difficulty(qpts), 1)
            pdf.cell(16, 7, str(qpts), 1)
            pdf.cell(24, 7, str(weight(q)), 1)
            pdf.ln()

        pdf.add_page()
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 8, 'Đáp án', ln=True)
        pdf.set_font("Arial", '', 12)
        for idx, q in enumerate(questions, 1):
            key = q.get("answer", "")
            pdf.multi_cell(0, 7, f"Câu {idx}: {key}")

        out_path = os.path.join("outputs", f"{safe_title}_{int(time.time())}.pdf")
        pdf.output(out_path)
        return send_file(out_path, as_attachment=True, download_name=os.path.basename(out_path))

    else:
        return jsonify({"error": "Định dạng không hỗ trợ"}), 400


@utilities_bp.route("/api/archive", methods=["POST"])
def archive_files():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "Không có file để nén"}), 400
    os.makedirs("outputs", exist_ok=True)
    zip_path = os.path.join("outputs", f"archive_{int(time.time())}.zip")
    with zipfile.ZipFile(zip_path, "w") as zf:
        for f in files:
            filename = f.filename
            if not filename:
                continue
            tmp_path = os.path.join("outputs", f"tmp_{int(time.time())}_{filename}")
            f.save(tmp_path)
            zf.write(tmp_path, filename)
            os.remove(tmp_path)
    return send_file(zip_path, as_attachment=True, download_name=os.path.basename(zip_path))


@utilities_bp.route("/api/unarchive", methods=["POST"])
def unarchive_file():
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    zip_file = request.files.get("file")
    if not zip_file:
        return jsonify({"error": "Không có file zip"}), 400
    os.makedirs("outputs", exist_ok=True)
    tmp_zip = os.path.join("outputs", f"tmp_unarchive_{int(time.time())}.zip")
    zip_file.save(tmp_zip)
    dest = os.path.join("outputs", f"unpacked_{int(time.time())}")
    os.makedirs(dest, exist_ok=True)
    with zipfile.ZipFile(tmp_zip, "r") as zf:
        zf.extractall(dest)
        files = zf.namelist()
    os.remove(tmp_zip)
    return jsonify({"path": dest, "files": files})


# ──────────────────────────────────────────
#  600+ TOOLS API ENDPOINTS
# ──────────────────────────────────────────

@utilities_bp.route("/api/tools/calc", methods=["POST"])
def tool_calc():
    data = request.get_json()
    return jsonify({"result": calculate_math(data.get("expression", ""))})


@utilities_bp.route("/api/tools/convert", methods=["POST"])
def tool_convert():
    data = request.get_json()
    return jsonify({"result": unit_converter(data.get("value", 0), data.get("from", ""), data.get("to", ""))})


@utilities_bp.route("/api/tools/currency", methods=["POST"])
def tool_currency():
    data = request.get_json()
    return jsonify({"result": currency_converter(data.get("amount", 0), data.get("from", "USD"), data.get("to", "VND"))})


@utilities_bp.route("/api/tools/bmi", methods=["POST"])
def tool_bmi():
    data = request.get_json()
    return jsonify({"result": bmi_calculator(data.get("weight", 60), data.get("height", 170))})


@utilities_bp.route("/api/tools/password", methods=["POST"])
def tool_password():
    data = request.get_json()
    return jsonify({"result": generate_password(data.get("length", 16), data.get("symbols", True))})


@utilities_bp.route("/api/tools/qr", methods=["POST"])
def tool_qr():
    data = request.get_json()
    text = data.get("text", "")
    import urllib.parse
    encoded = urllib.parse.quote(text)
    url = f"https://api.qrserver.com/v1/create-qr-code/?size=300x300&data={encoded}"
    return jsonify({"result": qr_code_url(text), "image_url": url, "text": text})


@utilities_bp.route("/api/tools/hash", methods=["POST"])
def tool_hash():
    data = request.get_json()
    return jsonify({"result": hash_text(data.get("text", ""), data.get("algo", "sha256"))})


@utilities_bp.route("/api/tools/base64", methods=["POST"])
def tool_base64():
    data = request.get_json()
    return jsonify({"result": base64_encode_decode(data.get("text", ""), data.get("mode", "encode"))})


@utilities_bp.route("/api/tools/countdown", methods=["POST"])
def tool_countdown():
    data = request.get_json()
    return jsonify({"result": countdown_timer(data.get("date", ""))})


@utilities_bp.route("/api/tools/random", methods=["POST"])
def tool_random():
    data = request.get_json()
    return jsonify({"result": random_number(data.get("min", 1), data.get("max", 100))})


@utilities_bp.route("/api/tools/color", methods=["POST"])
def tool_color():
    data = request.get_json()
    return jsonify({"result": color_picker_info(data.get("hex", "#7c5cff"))})


@utilities_bp.route("/api/tools/loan", methods=["POST"])
def tool_loan():
    data = request.get_json()
    return jsonify({"result": loan_calculator(data.get("principal", 0), data.get("rate", 0), data.get("months", 12))})


@utilities_bp.route("/api/tools/ip", methods=["POST"])
def tool_ip():
    data = request.get_json()
    ip = data.get("ip", "")
    result = ip_lookup(ip) if ip else my_ip()
    return jsonify({"result": result})


@utilities_bp.route("/api/tools/wordcount", methods=["POST"])
def tool_wordcount():
    data = request.get_json()
    return jsonify({"result": word_counter(data.get("text", ""))})


@utilities_bp.route("/api/tools/age", methods=["POST"])
def tool_age():
    data = request.get_json()
    return jsonify({"result": age_calculator(data.get("date", ""))})


@utilities_bp.route("/api/tools/translate", methods=["POST"])
def tool_translate():
    data = request.get_json()
    return jsonify({"result": translate_text(data.get("text", ""), data.get("lang", "vi"))})


@utilities_bp.route("/api/tools/stock", methods=["POST"])
def tool_stock():
    data = request.get_json()
    return jsonify({"result": get_stock_info(data.get("symbol", "AAPL"))})


@utilities_bp.route("/api/tools/crypto", methods=["POST"])
def tool_crypto():
    data = request.get_json()
    return jsonify({"result": get_crypto_price(data.get("coin", "bitcoin"))})


@utilities_bp.route("/api/tools/pomodoro", methods=["POST"])
def tool_pomodoro():
    data = request.get_json()
    return jsonify({"result": pomodoro_plan(data.get("work", 25), data.get("sessions", 4))})


@utilities_bp.route("/api/tools/lorem", methods=["POST"])
def tool_lorem():
    data = request.get_json()
    return jsonify({"result": generate_lorem_ipsum(data.get("paragraphs", 2))})


@utilities_bp.route("/api/tools/quote", methods=["GET"])
def tool_quote():
    return jsonify({"result": motivational_quote()})


@utilities_bp.route("/api/tools/define", methods=["POST"])
def tool_define():
    data = request.get_json()
    return jsonify({"result": define_word(data.get("word", ""))})


@utilities_bp.route("/api/tools/ping", methods=["POST"])
def tool_ping():
    data = request.get_json()
    return jsonify({"result": ping_website(data.get("url", ""))})


@utilities_bp.route("/api/tools/regex", methods=["POST"])
def tool_regex():
    data = request.get_json()
    return jsonify({"result": regex_tester(data.get("pattern", ""), data.get("text", ""))})


@utilities_bp.route("/api/tools/fibonacci", methods=["POST"])
def tool_fibonacci():
    data = request.get_json()
    return jsonify({"result": fibonacci(data.get("n", 10))})


@utilities_bp.route("/api/tools/prime", methods=["POST"])
def tool_prime():
    data = request.get_json()
    return jsonify({"result": prime_check(data.get("n", 7))})


@utilities_bp.route("/api/tools/binary", methods=["POST"])
def tool_binary():
    data = request.get_json()
    return jsonify({"result": binary_converter(data.get("value", "0"), data.get("from", "decimal"))})


@utilities_bp.route("/api/tools/caesar", methods=["POST"])
def tool_caesar():
    data = request.get_json()
    return jsonify({"result": caesar_cipher(data.get("text", ""), data.get("shift", 3), data.get("decode", False))})


@utilities_bp.route("/api/tools/palindrome", methods=["POST"])
def tool_palindrome():
    data = request.get_json()
    return jsonify({"result": check_palindrome(data.get("text", ""))})


@utilities_bp.route("/api/tools/time", methods=["GET"])
def tool_time():
    return jsonify({"result": get_current_time()})


@utilities_bp.route("/api/tools/palette", methods=["POST"])
def tool_palette():
    data = request.get_json()
    return jsonify({"result": color_palette_generate(data.get("color", "#7c5cff"))})


@utilities_bp.route("/api/tools/percent", methods=["POST"])
def tool_percent():
    data = request.get_json()
    return jsonify({"result": percentage_calc(data.get("value", 0), data.get("percent", 0), data.get("mode", "of"))})


@utilities_bp.route("/api/tools/studyplan", methods=["POST"])
def tool_studyplan():
    data = request.get_json()
    return jsonify({"result": study_plan(data.get("subject", ""), data.get("days", 30), data.get("hours", 2))})


# ──────────────────────────────────────────
#  10 NEW SOPHISTICATED TOOLS
# ──────────────────────────────────────────

@utilities_bp.route("/api/tools/seo", methods=["POST"])
def tool_seo():
    """SEO分析工具 - Analyze webpage SEO."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    return jsonify({"result": analyze_seo_url(data.get("url", ""))})


@utilities_bp.route("/api/tools/json", methods=["POST"])
def tool_json():
    """JSON validator and formatter."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    json_str = request.form.get("json", "") or data.get("json", "")
    return jsonify({"result": validate_json_format(json_str)})


@utilities_bp.route("/api/tools/email", methods=["POST"])
def tool_email():
    """Advanced email validation."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    return jsonify({"result": validate_email_advanced(data.get("email", ""))})


@utilities_bp.route("/api/tools/codequality", methods=["POST"])
def tool_codequality():
    """Code quality analyzer."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    code = request.form.get("code", "") or data.get("code", "")
    return jsonify({"result": analyze_code_quality(code)})


@utilities_bp.route("/api/tools/sentiment", methods=["POST"])
def tool_sentiment():
    """Text sentiment analysis."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    text = request.form.get("text", "") or data.get("text", "")
    return jsonify({"result": analyze_sentiment(text)})


@utilities_bp.route("/api/tools/urlinfo", methods=["POST"])
def tool_urlinfo():
    """Shortened URL analyzer."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    return jsonify({"result": compress_url_analyzer(data.get("url", ""))})


@utilities_bp.route("/api/tools/textsim", methods=["POST"])
def tool_textsim():
    """Text similarity comparison."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    text1 = request.form.get("text1", "") or data.get("text1", "")
    text2 = request.form.get("text2", "") or data.get("text2", "")
    return jsonify({"result": text_similarity_compare(text1, text2)})


@utilities_bp.route("/api/tools/html2text", methods=["POST"])
def tool_html2text():
    """HTML to text analyzer."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    html_input = request.form.get("html", "") or data.get("html", "")
    return jsonify({"result": html_to_text_analyzer(html_input)})


@utilities_bp.route("/api/tools/urlcheck", methods=["POST"])
def tool_urlcheck():
    """Detailed URL accessibility check."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    return jsonify({"result": check_url_accessibility(data.get("url", ""))})


@utilities_bp.route("/api/tools/sqlanalyze", methods=["POST"])
def tool_sqlanalyze():
    """SQL query analyzer."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    sql = request.form.get("sql", "") or data.get("sql", "")
    return jsonify({"result": analyze_sql_query(sql)})


@utilities_bp.route("/api/tools/run", methods=["POST"])
def tool_run_generic():
    """Chạy bất kỳ tool nào theo tên."""
    if "user_id" not in session:
        return jsonify({"error": "Chưa đăng nhập"}), 401
    data = request.get_json()
    tool = data.get("tool", "")
    args = data.get("args", "")
    result = run_tool(tool, str(args))
    return jsonify({"result": result, "tool": tool})